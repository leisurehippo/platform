# -*- coding: utf-8 -*- 
import jieba
import jieba.posseg as posseg
import re
import os
from docx import Document
import codecs
import traceback
import argparse
import chardet
manufacturer_dict = set()
device_type_dict = set()
parts_dict=set()
location_dict=set()
level_dict=set()

equip_pattern=re.compile(u"(\d+[k千]*[伏v]+)([^-线站厂所,.，。；;、：:(（）)]*[线站厂所])",re.I)
ta_pattern=re.compile(u'#(\d+)')
size_pattern=re.compile(u'([\w\d]+[-/*×.()]+){1,10}[\w\d]+\)*',re.I)
distance_pattern=re.compile(u'\d+\.*\d+k*m|^\d+\.\d+$',re.I)
company_stop_words=[u'山东送变电工程公司',u'LG']
runtime_pat=re.compile(u'\d+年(\d+月)*(\d+日)*|\d+-\d+-\d+')


def get_substation(substation, line):
    if len(substation) > 0:
        return substation
    line = line.replace(' ','').strip()
    line = re.sub(u'([kK][vV])|(千伏)','kV',line)
    p_substation = u'.+变电站'
    m3 = re.search(p_substation, line)
    if m3 != None and len(substation)==0:
        seg_list = list(jieba.cut(m3.group(0).replace(' ','')))
        res = ''.join(seg_list[-3:])
        if u'kV' in res:
            if len(seg_list[-1]) != 3:
                substation = seg_list[-1]
            else:
                substation = re.sub(u'\d+kV','',res)
        else:
            substation = ''.join(seg_list[-2:])
    block_words = [u'在', u'于', u'发现', u'对', u'该', u'×', u'某', u'所有',u'X',u'此',u'为']
    for word in block_words:
        if word in substation:
            substation = u'某变电站'
            break

    return substation



def get_unit_name(unit_name, line):
    if u'单位名称' in line and len(unit_name) == 0:
        line = line.replace(u'：', ':')
        words = line.split(u':')
        if len(words)>1:
            unit_name= words[1]
        else:
            unit_name = line[4:]
        unit_name = unit_name.strip().strip(u'【').strip(u'】')
    return unit_name

def get_device(device_type, line):
    if len(device_type) != 0:
        return device_type
    for device in device_type_dict:
        #device = device.encode('UTF-8')
        if device in line:
            device_type = device
            break
    return device_type

def get_xianlu(device_type,line):
    if len(device_type) != 0:
        return device_type
    for m in equip_pattern.finditer(line):
        device_type=m.group(2)
        if len(device_type)>1:
            break
    return device_type

def get_voltage(voltage, line):
    line = line.replace(' ','')
    line = re.sub(u'([kK][vV])|(千伏)','kV',line)
    m = re.search(u'\d+kV', line)
    if m != None and len(voltage) == 0:
        voltage = m.group(0)
    return voltage

def get_run_date(run_date, line):
    if len(run_date) != 0:
        return run_date
    line = line.replace(u'：',u':').replace(' ','')
    date = u'\d+年(\d+月)*(\d+日?)*'
    p1 = date + u'((投入运行)|(投运))'
    p2 = u'投运日期为' + date
    p3 = u'投运日期为'+u'\d+-\d+-\d+'
    p4 = u'投运日期:' + date
    # m1 = re.search(p1, line)
    # m2 = re.search(p2, line)
    # m3 = re.search(p3, line)
    # m4 = re.search(p4, line)
    if u'投运' in line:
        m5=re.search(runtime_pat,line)
        if m5:
            run_date=m5.group(0)

    if not run_date:
        m1 = re.search(p1, line)
        if m1 != None:
            run_date = m1.group(0).replace(u'投入运行',u'投运')[:-2]
    # elif m2 != None:
    # 	run_date = m2.group(0)[5:]
    # elif m3 != None:
    # 	run_date = m3.group(0)[5:]
    # elif m4 != None:
    # 	run_date = m4.group(0)[5:]

    return run_date.replace(u'年','/').replace(u'月','/').replace(u'日','').strip('/')


# ust find the first date appear in the head of line
def get_find_date(accident_date, line):
    if len(accident_date) != 0:
        return accident_date
    line = line.strip()
    for word in re.split(u'[。；]', line):
        p = u'^\d+年(\d+月)*(\d+日)*'
        m = re.search(p, word.replace(' ',''))
        if m != None:
            accident_date = m.group(0)
            break
    return accident_date

def get_device_id(device_id, line):
    if len(device_id) !=0 :
        return device_id

    m = re.search(u'([A-Za-z0-9]+[-/]+){1,10}[A-Z0-9\.]+', line)
    if m != None and re.search(u'[A-Z]+', m.group(0)) != None:
        device_id = m.group(0)
    return device_id

def get_device_id2(device_id,line):
    for li in re.split(u'[,，。；、：？:?!！]',line):
        for word in jieba.cut(li,HMM=False):
            if word in parts_dict:
                possible_size=size_pattern.search(li)
                #找到一对就行，一句话也不会说两对吧？
                if possible_size:
                    #会识别出(杆塔,22.45km)这种错误型号，过滤掉
                    if not distance_pattern.match(possible_size.group()):
                        device_id.add((word,possible_size.group()))
                        break
    return device_id


def statistic(cases):
    res = {}
    res['total'] = len(cases)
    for case in cases:
        for key in case.keys():
            name, val = case[key]
            if name not in res:
                res[name] = 0
            if len(val) != 0:
                res[name] += 1
    return res


def show_statistic(res):
    print('total：%d' % res['total'])
    for key in res:
        print(key + ':' + str(res[key]) + '/' + str(res['total']))

#可以作为提取生产厂家名称的规则，用来扩充词典
def get_manufacturer(manufacturer, line):
    if len(manufacturer) !=0 :
        return manufacturer
    for m in manufacturer_dict:
        if m in line:
            return m
    replace_words = [u'为',u'了', u'系', u'由', u'是', u'，', u'使用', u' ', u':', u'利用', u'结合']
    block_list = [u'系',u'员工',u'于',u'该']
    if u'生产' in line and len(manufacturer)==0:
        lines = re.split(u'[，,.。;；]',line)
        for l in lines:
            if u'生产' in l:
                line = l
                break
        line = line.replace(' ','').replace(u'：',' ').strip()
        for word in replace_words:
            line = line.replace(word, u'系')
        p2 = u'生产厂家系.*'
        m2 = re.search(p2, line)
        if m2 != None:
            manufacturer = m2.group(0)[5:]
            return manufacturer

        p1 = u'系.*(生产|制造)'
        m1 = re.search(p1, line)
        if m1 != None:
            manufacturer = m1.group(0)[1:-2]
        if len(line) > 2 and line[-2:] == u'生产' and len(manufacturer)==0:
            manufacturer = line[:-2]
    manufacturer = re.sub(u'[于在]?\d+年(\d+月)?(\d+日?)?','',manufacturer)
    if len(manufacturer) > 0:
        for word in block_list:
            if word in manufacturer:
                manufacturer = ''
                break
    return manufacturer

def get_manufacturer2(manufacturer, line):
    file_words=[]
    for word in jieba.cut(line,HMM=False):
        file_words.append(word)
        if word in manufacturer_dict and word not in company_stop_words:
            #厂家是先找厂家再往前找设备，因为基本都是先说设备再说厂家，但型号可能在设备前也可能在后面
            #所以找一对就算了
            for i in range(-1,-len(file_words)-1,-1):
                if file_words[i] in parts_dict:
                    manufacturer.add((file_words[i],word))
                    break

    return manufacturer

def get_weather(weather, line):
    if len(weather) != 0:
        return weather
    for word in line.split(' '):
        if u'天气：' in word:
            word = word.replace(u'情况','')
            weather = word.split(u'：')[-1].replace(u'。','')
    return weather

def get_weather_2(weather, line):
    if len(weather) != 0:
        return weather
    res = []
    weather_key = [u'雨', u'雪', u'雾', u'阴', u'晴', u'多云']
    for weather_word in weather_key:
        if weather_word in line and weather_word not in set(weather):
            res.append(weather_word)
    if len(res) == 2:
        weather = res[0] + u'转' + res[1]
    elif len(res) > 2:
        weather = u'晴'
    else:
        weather = res[0]
    return weather
def get_temperature(temperature, line):
    if len(temperature) != 0:
        return temperature
    if u'温度' in line:
        line = line.replace(' ','')
        for word in re.split(u'[,。，；]',line):
            if u'测试温度' in word or u'环境温度' in word or u'测试环境' in word:
                m = re.search(u'-?\d+[度(℃)(°C)]',word)
                if m != None:
                    temperature = re.search(u'-?\d+',m.group(0)).group(0) + u'℃'
    return temperature


def get_humidity(humidity, line):
    line = line.replace(' ','')
    if len(humidity) != 0:
        return humidity
    for word in re.split(u'[，。；]', line):
        if u'湿度' in word:
            m = re.search(u'\d+[%％]', word)
            if m != None:
                humidity = m.group(0)
    return humidity

def get_fault_1(file_name):
    file_name = file_name.replace(u'-','')
    res = ''
    detect_words = [u'发现',u'检测',u'诊断',u'检测']
    for word in detect_words:
        if word in file_name:
            res = file_name.split(u'发现')[-1]
            break
    return res


def get_fault_2(file_name):
    file_name = file_name.strip().replace('-',' ').replace('_',' ')
    res = ''.join(file_name.split(' ')[-3:])
    return res

def get_fault_xianlu(file_name):
    file_name = file_name.strip().replace('-',' ').replace('_',' ')
    file_name=file_name.strip().split(' ')
    res=''
    if len(file_name)>1:
        if file_name[-1].isdigit():
            res=file_name[-2]
        else:
            res=file_name[-1]
    return res

def deal_info(cases, key):
    res = []
    for i in range(len(cases)):
        cases[i].pop(key,None)

def get_location(location,line):
    for word in jieba.cut(line,HMM=False):
        if word in location_dict:
            location.add(word)
    return location

def get_badlevel(bad_level,line):
    if  bad_level:
        return bad_level
    for level in level_dict:
        if level in line:
            bad_level=level
    return bad_level

def get_tower_id(tower_id,line):
    if tower_id:
        return tower_id
    for m in ta_pattern.finditer(line):
        tower_id=set([(u'塔号',m.group(1))])
    return tower_id

def extract_info(root, xianlu=False):
    dir_path = root + '/newdir'
    docx_path = root + '/docx/'
    cases = []
    file_names = os.listdir(dir_path)
    for _filename in file_names:
        if _filename[-4:]!='.txt' or  '~'  in _filename or _filename =='fileNames.txt':
            continue
        # if _filename!='07-新疆-射频局部放电检测发现220kV变压器套管局部放电缺陷.txt':
        # 	continue
        file_path=dir_path+'/'+_filename
        try:
            file=codecs.open(file_path,'r')
            lines = file.readlines()
        except UnicodeDecodeError:
            file = codecs.open(file_path,'r',encoding='latin-1')
            lines = file.readlines()
        try:
            case_dict = analyse_table(docx_path+_filename[:-4]+'.docx')
        except IndexError as e:
            case_dict={}
            traceback.print_exc()
        file_name =  _filename[:-4].decode('GB2312')     
        case_dict[u'报告名称']=(u'报告名称',file_name)
        if not xianlu:
            fault = get_fault_1(file_name)
        else:
            fault=get_fault_xianlu(file_name)
        case_dict[u'故障名称']=(u'故障名称', fault)
        device_type = ''
        unit_name = ''
        substation = ''
        voltage = ''
        voltage = get_voltage(voltage, file_name)
        run_date = ''
        device_id = set()
        manufacturer = set()
        location=set()
        other=set()
        tower_id=set()
        temperature = ''
        humidity = ''
        find_date = ''
        weather = ''
        bad_level=''
        for i in range(len(lines)):
            line = lines[i].strip()
            line = line.decode('gb2312')
            if '\r' in line or len(line)<2:
                continue
            line=re.sub('[\s\u3000]+','',line)
            unit_name = get_unit_name(unit_name, line)
            substation = get_substation(substation, line)
            voltage = get_voltage(voltage, line)
            run_date = get_run_date(run_date, line)
            device_id = get_device_id2(device_id, line)
            manufacturer = get_manufacturer2(manufacturer, line)
            temperature = get_temperature(temperature, line)
            humidity = get_humidity(humidity, line)
            find_date = get_find_date(find_date, line)
            weather = get_weather(weather, line)
            location=get_location(location,line)
            bad_level=get_badlevel(bad_level,line)

            if not xianlu:
                device_type = get_device(device_type, line)
            else:
                device_type=get_xianlu(device_type,line)
                tower_id=get_tower_id(tower_id,line)

        if substation == u'某变电站':
            substation = ''
        case_dict[u'单位名称']=(u'单位名称', unit_name)
        case_dict[u'变电站名称']=(u'变电站名称', substation)
        case_dict[u'电压等级']=(u'电压等级', voltage)
        case_dict[u'投运时间']=(u'投运时间', run_date)
        if u'型号'  in case_dict:
            device_id|=case_dict[u'型号'][1]
        case_dict[u'型号']=('型号', device_id)
        case_dict[u'生产厂家']=('生产厂家', manufacturer)
        case_dict[u'设备类型']=('设备类型', device_type)
        case_dict[u'测试环境温度']=('测试环境温度', temperature)
        case_dict[u'测试环境湿度']=('测试环境湿度', humidity)
        case_dict[u'故障发现时间']=('故障发现时间', find_date)
        case_dict[u'测试时天气']=('测试时天气', weather)
        case_dict[u'缺陷等级']=('缺陷等级',bad_level)
        case_dict[u'地理位置']=('地理位置',location)
        if xianlu:
            if u'塔号' in case_dict:
                tower_id|=case_dict[u'塔号'][1]
            case_dict[u'塔号']=(u'塔号',tower_id)
        if len(fault) == 0 and len(device_type)>0:
            case_dict[u'故障名称'] = (u'故障名称',device_type+u'故障')
        cases.append(case_dict)
    return cases


def statistic(cases):
    res = {}
    res['total'] = len(cases)
    for case in cases:
        for key in case:
            name,val=case[key]
            if name not in res:
                res[name] = 0
            if len(val) != 0:
                res[name] += 1
    return res


def show_statistic(res):
    print('total：%d' % res['total'])
    for key in res:
        print(key + ':' + str(res[key]) + '/' + str(res['total']))

def test_table(filename):
    doc=Document(filename)
    print(len(doc.tables))
    rows=doc.tables[0].rows
    print(len(rows))
    for row in rows:
        for cell in row.cells:
            print(cell.text,'---->>')
        print('#########')

def analyse_table(fullfilename):
    key_words=[u'塔号',u'型号',u'时间',u'故障前',u'故障时',u'保护动作',u'保护装置']
    doc=Document(fullfilename)
    # test_table(fullfilename)
    # return result_dict
    case_dict={}
    result_dict={u'塔号':set(),u'型号':set(),u'其他时间':set(),u'故障前情况':set(),u'故障时情况':set(),u'投运时间':''\
    ,u'保护动作':set(),u'保护装置':set()}
    for t in doc.tables:
        lastrow=[]
        last_valid_index={}
        for row in t.rows:
            currow=[]
            valid_index={}
            for ind,cell in enumerate(row.cells):
                cell=cell.text.strip()
                currow.append(cell)
                for key in key_words:
                    if key in cell:#查找含有关键词的表头
                        if key==u'故障前' or key==u'故障时':
                            key+=u'情况'
                        if key==u'时间' :
                            if cell==u'投运时间':
                                key=cell
                            else:
                                key=u'其他时间'
                        valid_index[key]=valid_index.get(key,[])
                        valid_index[key].append(ind)
                        break

            if len(set(currow)&set(lastrow))>1:#相邻两行有交集的话，说明是复杂表头被拆分成多行的原因
                lastrow=currow
                last_valid_index=valid_index
                continue
            else:
                if lastrow:#lastrow为空表示这是表头行不是内容行
                    for key in last_valid_index:
                        #result_dict[key]=result_dict.get(key,set())
                        for index in last_valid_index[key]:
                            if currow[index] and lastrow[index]:
                                if key!=u'投运时间':
                                    result_dict[key].add((lastrow[index],currow[index]))
                                else:
                                    result_dict[key]=currow[index]
                    lastrow=[]
                    last_valid_index={}
                else:
                    lastrow=currow
                    last_valid_index=valid_index

    for key in result_dict:
        case_dict[key]=(key,result_dict[key])
    return case_dict


"""
columns = ["file_name","fault_name","substation","tower_id","voltaget","device_type","device_id","manufacturer",
"run_date","find_date","weather","temperature","humidity",
"location","bad_level","protective_action","protective_device","situation","other_time"]
"""
def save_to_file(cases, file_path):
    #write to target file
    columns = [u'报告名称',u'故障名称', u'单位名称', u'变电站名称', u'塔号',u'电压等级', u'设备类型',u'型号',u'生产厂家',
    u'投运时间', u'故障发现时间',u'测试时天气',u'测试环境温度', u'测试环境湿度',
    u'地理位置',u'缺陷等级', u'保护动作',u'保护装置',u'故障前情况',u'故障时情况',u'其他时间']
    result = codecs.open(file_path, 'w', encoding='utf8')
    for case in cases:
        for key in columns:
            if len(case[key][1]) == 0:
                val = '#'
            else:
                val = case[key][1]

            if type(val) == set:
                val = ','.join([v[0] +':'+ v[1] if type(v) == tuple else v for v in val])
            val = val.replace(' ','').replace('\n','')
            line = val +' '
            result.write(line)
        result.write('\n')
    result.close()
    
def main(root_dir, dict_path, xianlu):
    print("main....")
    manufacturer_filename = dict_path +'/merge_company.txt'
    device_type_file_name = dict_path +'/device_type.txt'
    parts_filename=dict_path +'/parts.txt'
    location_filename=dict_path +'/region.txt'
    level_filename=dict_path +'/level.txt'
    user_dictfile= dict_path +'/user_dict.txt'
    equipment_filename=dict_path +'/equipment.txt'
    errors_filename=dict_path +'/error.txt'
    loc_filename=dict_path +'/location.txt'

    print("loading...")
    jieba.load_userdict(user_dictfile)
    jieba.load_userdict(manufacturer_filename)
    jieba.load_userdict(equipment_filename)
    jieba.load_userdict(errors_filename)
    jieba.load_userdict(parts_filename)
    jieba.load_userdict(level_filename)
    jieba.load_userdict(loc_filename)
    jieba.load_userdict(location_filename)
    jieba.load_userdict(device_type_file_name)

    manufacturer_file=codecs.open(manufacturer_filename,'r',encoding='utf-8')
    for line in manufacturer_file:
        manufacturer_dict.add(line.strip())

    with codecs.open(device_type_file_name, 'r', encoding='utf-8') as f:
        for line in f.readlines():
            device_type_dict.add(line.strip())

    with codecs.open(parts_filename,'r',encoding='utf-8') as f:
        for line in f:
            parts_dict.add(line.strip())

    with codecs.open(location_filename,'r',encoding='utf-8') as f:
        for line in f:
            location_dict.add(line.strip())

    with codecs.open(level_filename,'r',encoding='utf-8') as f:
        for line in f:
            level_dict.add(line.strip())

    print("begin....")
    cases = extract_info(root_dir,xianlu)
    save_to_file(cases, root_dir+'/'+'result.txt')
    print("end")

if __name__ == '__main__':
    import sys
    print("running...")
    if len(sys.argv)>4 or len(sys.argv)<=1:
        print('参数数目异常!')
    else:
        if len(sys.argv)==3:
            print(sys.argv[1])
            main(sys.argv[1],sys.argv[2],True)
        else:
            main(sys.argv[1],sys.argv[2],False)